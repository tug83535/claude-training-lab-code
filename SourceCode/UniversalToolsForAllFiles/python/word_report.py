"""
KBT Universal Tools — Word Report Generator
============================================================
Reads analytical outputs from Excel and generates a formatted
Microsoft Word document automatically — tables, headings, and all.

Usage:
    python word_report.py "C:\\path\\data.xlsx"
    python word_report.py "data.xlsx" --title "Q1 Finance Report" --author "Connor"
    python word_report.py "data.xlsx" --sheets "Summary" "Detail"

Output: Saves "WORD_REPORT.docx" in the same folder as the Excel file
"""

import sys
import os
import argparse
from datetime import datetime

try:
    import pandas as pd
    import openpyxl
except ImportError:
    print("ERROR: Run: pip install pandas openpyxl")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed.")
    print("Run: pip install python-docx")
    sys.exit(1)


def generate_word_report(file_path: str, title: str, author: str,
                          sheets: list = None) -> None:
    print(f"\n{'='*55}")
    print("  KBT Word Report Generator")
    print(f"{'='*55}")
    print(f"  File:   {os.path.basename(file_path)}")
    print(f"  Title:  {title}")
    print(f"  Author: {author}")
    print(f"  Date:   {datetime.now().strftime('%m/%d/%Y %I:%M %p')}")
    print(f"{'='*55}\n")

    if not os.path.exists(file_path):
        print(f"ERROR: File not found: {file_path}")
        sys.exit(1)

    xl = pd.ExcelFile(file_path)
    sheets_to_include = sheets if sheets else xl.sheet_names
    invalid = [s for s in sheets_to_include if s not in xl.sheet_names]
    if invalid:
        print(f"ERROR: Sheet(s) not found: {invalid}")
        print(f"Available: {xl.sheet_names}")
        sys.exit(1)

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"Prepared by: {author}     |     Date: {datetime.now().strftime('%B %d, %Y')}")
    sub_run.font.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()

    for sheet_name in sheets_to_include:
        print(f"  Processing sheet: {sheet_name}")
        try:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            df = df.fillna("")
        except Exception as e:
            print(f"  ERROR reading '{sheet_name}': {e}")
            continue

        # Section heading
        doc.add_heading(sheet_name, level=1)

        # Brief summary paragraph
        summary = doc.add_paragraph()
        summary_run = summary.add_run(
            f"This section contains data from the '{sheet_name}' sheet. "
            f"Total records: {len(df):,}. Columns: {len(df.columns)}."
        )
        summary_run.font.size = Pt(11)

        # Table (cap at 50 rows to keep document reasonable)
        display_df = df.head(50)
        if len(df) > 50:
            note = doc.add_paragraph()
            note.add_run(f"Note: Showing first 50 of {len(df):,} rows. See the Excel file for complete data.").italic = True

        table = doc.add_table(rows=1, cols=len(display_df.columns))
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(display_df.columns):
            cell = header_cells[col_idx]
            cell.text = str(col_name)
            # Style header cell
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            # Set header background color
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = tcPr.find(qn('w:shd'))
            if shd is None:
                from docx.oxml import OxmlElement
                shd = OxmlElement('w:shd')
                tcPr.append(shd)
            shd.set(qn('w:fill'), '1F497D')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')

        # Data rows
        for _, row_data in display_df.iterrows():
            row_cells = table.add_row().cells
            for col_idx, val in enumerate(row_data):
                row_cells[col_idx].text = str(val)
                for para in row_cells[col_idx].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

    # Footer note
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"Generated by KBT Universal Tools on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}. "
        f"Source file: {os.path.basename(file_path)}."
    )
    footer_run.font.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = os.path.join(os.path.dirname(file_path), "WORD_REPORT.docx")
    doc.save(output_path)

    print(f"\n{'='*55}")
    print(f"  DONE! Word document saved to:")
    print(f"  {output_path}")
    print(f"{'='*55}\n")


def main():
    parser = argparse.ArgumentParser(description='KBT Word Report Generator')
    parser.add_argument('file', help='Path to the Excel file to convert')
    parser.add_argument('--title', default='Finance Report',
                        help='Report title (default: "Finance Report")')
    parser.add_argument('--author', default='Finance Team',
                        help='Author name for the report (default: "Finance Team")')
    parser.add_argument('--sheets', nargs='+', default=None,
                        help='Specific sheet names to include (default: all sheets)')
    args = parser.parse_args()
    generate_word_report(args.file, args.title, args.author, args.sheets)


if __name__ == '__main__':
    main()
