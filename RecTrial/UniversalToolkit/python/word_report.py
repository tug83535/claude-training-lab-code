"""
KBT Universal Tools — Word Report Generator
============================================================
Reads analytical outputs from Excel and generates a formatted
Microsoft Word document automatically — tables, headings, and all.

Usage:
    python word_report.py "C:\\path\\data.xlsx"
    python word_report.py "data.xlsx" --title "Q1 Finance Report" --author "Connor"
    python word_report.py "data.xlsx" --sheets "Summary" "Detail"

Output: Saves "WORD_REPORT.docx" in the same folder as the Excel file
"""

import sys
import os
import argparse
from datetime import datetime

try:
    import pandas as pd
    import openpyxl
except ImportError:
    print("ERROR: Run: pip install pandas openpyxl")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed.")
    print("Run: pip install python-docx")
    sys.exit(1)


def build_talking_points(dfs_by_sheet: dict) -> list:
    """Auto-generate 3-5 plain-English talking points from the report data.

    Cherry-picked concept from Codex comparison (Batch 3, 2026-04-21).
    Takes a dict of {sheet_name: DataFrame} and returns a list of short
    sentences suitable for a CFO one-pager.
    """
    points = []
    if not dfs_by_sheet:
        return points

    # Totals across all sheets
    total_rows = sum(len(df) for df in dfs_by_sheet.values())
    total_cells = sum(df.size for df in dfs_by_sheet.values())
    points.append(
        f"Report covers {len(dfs_by_sheet)} sheet(s) with {total_rows:,} total data rows "
        f"across {total_cells:,} cells."
    )

    # Largest sheet by row count
    largest = max(dfs_by_sheet.items(), key=lambda kv: len(kv[1]))
    points.append(
        f"Largest dataset: '{largest[0]}' with {len(largest[1]):,} rows "
        f"and {len(largest[1].columns)} columns."
    )

    # Data completeness (non-blank %)
    non_blank = sum((df != "").sum().sum() for df in dfs_by_sheet.values())
    completeness = (non_blank / total_cells * 100) if total_cells else 0
    points.append(f"Data completeness: {completeness:.1f}% of cells contain values.")

    # Sum of any "Amount" / "Total" / "Revenue" column we can find
    finance_keywords = ("amount", "total", "revenue", "actual", "expense")
    for sheet_name, df in dfs_by_sheet.items():
        for col in df.columns:
            if any(kw in str(col).lower() for kw in finance_keywords):
                try:
                    numeric = pd.to_numeric(df[col], errors="coerce").dropna()
                    if len(numeric) > 0:
                        total = numeric.sum()
                        points.append(
                            f"Total '{col}' across '{sheet_name}': "
                            f"${total:,.0f} (n={len(numeric):,})."
                        )
                        break
                except Exception:
                    continue
        if len(points) >= 4:
            break

    # Numeric column count (as a data-shape hint)
    numeric_col_total = 0
    for df in dfs_by_sheet.values():
        for col in df.columns:
            try:
                if pd.to_numeric(df[col], errors="coerce").notna().sum() > 0:
                    numeric_col_total += 1
            except Exception:
                continue
    if numeric_col_total > 0:
        points.append(
            f"Numeric columns available for analysis: {numeric_col_total} "
            f"across all sheets."
        )

    return points[:5]  # Cap at 5 points


def generate_word_report(file_path: str, title: str, author: str,
                          sheets: list = None, talking_points: bool = False) -> None:
    print(f"\n{'='*55}")
    print("  KBT Word Report Generator")
    print(f"{'='*55}")
    print(f"  File:   {os.path.basename(file_path)}")
    print(f"  Title:  {title}")
    print(f"  Author: {author}")
    print(f"  Date:   {datetime.now().strftime('%m/%d/%Y %I:%M %p')}")
    print(f"{'='*55}\n")

    if not os.path.exists(file_path):
        print(f"ERROR: File not found: {file_path}")
        sys.exit(1)

    xl = pd.ExcelFile(file_path)
    sheets_to_include = sheets if sheets else xl.sheet_names
    invalid = [s for s in sheets_to_include if s not in xl.sheet_names]
    if invalid:
        print(f"ERROR: Sheet(s) not found: {invalid}")
        print(f"Available: {xl.sheet_names}")
        sys.exit(1)

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"Prepared by: {author}     |     Date: {datetime.now().strftime('%B %d, %Y')}")
    sub_run.font.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()

    # Accumulate DataFrames for optional talking-points section
    dfs_by_sheet = {}

    for sheet_name in sheets_to_include:
        print(f"  Processing sheet: {sheet_name}")
        try:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            df = df.fillna("")
            dfs_by_sheet[sheet_name] = df
        except Exception as e:
            print(f"  ERROR reading '{sheet_name}': {e}")
            continue

        # Section heading
        doc.add_heading(sheet_name, level=1)

        # Brief summary paragraph
        summary = doc.add_paragraph()
        summary_run = summary.add_run(
            f"This section contains data from the '{sheet_name}' sheet. "
            f"Total records: {len(df):,}. Columns: {len(df.columns)}."
        )
        summary_run.font.size = Pt(11)

        # Table (cap at 50 rows to keep document reasonable)
        display_df = df.head(50)
        if len(df) > 50:
            note = doc.add_paragraph()
            note.add_run(f"Note: Showing first 50 of {len(df):,} rows. See the Excel file for complete data.").italic = True

        table = doc.add_table(rows=1, cols=len(display_df.columns))
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(display_df.columns):
            cell = header_cells[col_idx]
            cell.text = str(col_name)
            # Style header cell
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            # Set header background color
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = tcPr.find(qn('w:shd'))
            if shd is None:
                from docx.oxml import OxmlElement
                shd = OxmlElement('w:shd')
                tcPr.append(shd)
            shd.set(qn('w:fill'), '1F497D')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')

        # Data rows
        for _, row_data in display_df.iterrows():
            row_cells = table.add_row().cells
            for col_idx, val in enumerate(row_data):
                row_cells[col_idx].text = str(val)
                for para in row_cells[col_idx].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

    # Optional "Suggested Talking Points" section (opt-in via --talking-points)
    if talking_points and dfs_by_sheet:
        print("  Building suggested talking points...")
        doc.add_heading("Suggested Talking Points", level=1)
        intro = doc.add_paragraph()
        intro_run = intro.add_run(
            "Plain-English highlights auto-generated from the data above. "
            "Review and adapt for your audience before sharing."
        )
        intro_run.font.italic = True
        intro_run.font.size = Pt(10)
        intro_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        points = build_talking_points(dfs_by_sheet)
        for point in points:
            bullet_para = doc.add_paragraph(style="List Bullet")
            run = bullet_para.add_run(point)
            run.font.size = Pt(11)

        doc.add_paragraph()

    # Footer note
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"Generated by KBT Universal Tools on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}. "
        f"Source file: {os.path.basename(file_path)}."
    )
    footer_run.font.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = os.path.join(os.path.dirname(file_path), "WORD_REPORT.docx")
    doc.save(output_path)

    print(f"\n{'='*55}")
    print(f"  DONE! Word document saved to:")
    print(f"  {output_path}")
    print(f"{'='*55}\n")


def main():
    parser = argparse.ArgumentParser(description='KBT Word Report Generator')
    parser.add_argument('file', help='Path to the Excel file to convert')
    parser.add_argument('--title', default='Finance Report',
                        help='Report title (default: "Finance Report")')
    parser.add_argument('--author', default='Finance Team',
                        help='Author name for the report (default: "Finance Team")')
    parser.add_argument('--sheets', nargs='+', default=None,
                        help='Specific sheet names to include (default: all sheets)')
    parser.add_argument('--talking-points', action='store_true',
                        help='Append an auto-generated "Suggested Talking Points" section '
                             'with 3-5 plain-English highlights for CFO/exec audiences.')
    args = parser.parse_args()
    generate_word_report(args.file, args.title, args.author, args.sheets,
                         args.talking_points)


if __name__ == '__main__':
    main()
